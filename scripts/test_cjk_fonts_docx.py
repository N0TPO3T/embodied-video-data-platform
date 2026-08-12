from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

fonts = [
    "Hiragino Sans GB",
    "Hiragino Sans GB W3",
    "Heiti SC",
    "STHeiti",
    "STHeiti Medium",
    "Arial Unicode MS",
    "Arial Unicode",
    "Songti SC",
    "Apple LiGothic",
    "SimSun",
]

doc = Document()
for font in fonts:
    p = doc.add_paragraph()
    run = p.add_run(f"{font}: 具身智能视频数据平台 中文渲染测试")
    run.font.name = font
    run.font.size = Pt(16)
    run._element.get_or_add_rPr()
    r_fonts = run._element.rPr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font)

out = Path("outputs/019fcabc-e414-7b02-addb-e06f9e8ba4e3/cjk-font-test.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(out)
