from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "docs"
    / "superpowers"
    / "specs"
    / "2026-08-05-embodied-video-platform-requirements-architecture-design.md"
)
OUTPUT_DIR = ROOT / "outputs" / "019fcabc-e414-7b02-addb-e06f9e8ba4e3"
OUTPUT = OUTPUT_DIR / "具身智能视频数据平台-整体需求与技术实现方案-v1.0.docx"

NAVY = "183153"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D8DEE8"
WHITE = "FFFFFF"
RISK_RED = "9B1C1C"
GOLD = "A06A00"
POSITIVE = "1F5D42"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(
    run,
    *,
    latin="Arial Unicode MS",
    east_asia="Arial Unicode MS",
    size=None,
    color=None,
    bold=None,
    italic=None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr()
    r_fonts = run._element.rPr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, *, bottom=None, left=None) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name, spec in (("bottom", bottom), ("left", left)):
        if not spec:
            continue
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(spec.get("size", 8)))
        edge.set(qn("w:space"), str(spec.get("space", 2)))
        edge.set(qn("w:color"), spec.get("color", BLUE))
        p_bdr.append(edge)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("具身智能视频数据平台  |  整体需求与技术实现方案")
        set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        r1 = fp.add_run("第 ")
        set_run_font(r1, size=8.5, color=MUTED)
        add_page_field(fp)
        r2 = fp.add_run(" 页")
        set_run_font(r2, size=8.5, color=MUTED)


def set_style_font(style, size, color, bold=False) -> None:
    style.font.name = "Arial Unicode MS"
    style._element.get_or_add_rPr()
    r_fonts = style._element.rPr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        style._element.rPr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial Unicode MS")
    r_fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Callout" not in doc.styles:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Callout"]
    set_style_font(callout, 10.5, NAVY, False)
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2

    if "Diagram" not in doc.styles:
        diagram = doc.styles.add_style("Diagram", WD_STYLE_TYPE.PARAGRAPH)
    else:
        diagram = doc.styles["Diagram"]
    set_style_font(diagram, 9.5, NAVY, True)
    diagram.paragraph_format.left_indent = Inches(0.2)
    diagram.paragraph_format.right_indent = Inches(0.2)
    diagram.paragraph_format.space_before = Pt(5)
    diagram.paragraph_format.space_after = Pt(8)
    diagram.paragraph_format.line_spacing = 1.2


def add_numbering(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_inline(paragraph, text: str, *, size=None, color=None, bold=False) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                latin="Menlo",
                east_asia="Arial Unicode MS",
                size=(size or 10.5) - 0.3,
                color=DARK_BLUE,
                bold=False,
            )
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color, bold=bold)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("PRODUCT & TECHNICAL SPECIFICATION")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("具身智能视频数据平台")
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("整体需求清单与技术实现方案")
    set_run_font(run, size=17, color=DARK_BLUE, bold=False)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(26)
    set_paragraph_border(rule, bottom={"size": 12, "space": 1, "color": BLUE})

    meta_rows = [
        ("版本", "v1.0（技术沟通版）"),
        ("日期", "2026-08-05"),
        ("规划口径", "约 1000 名总用户｜Web 管理端 + 微信小程序成员端"),
        ("推荐架构", "D1 + R2 + Queues + 外部媒体/AI Worker"),
        ("使用场景", "需求评审、技术选型、开发拆分与一期验收"),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    set_table_geometry(table, [1800, 7560], indent_dxa=120)
    set_table_borders(table, color=WHITE, size=0)
    for row_index, (label, value) in enumerate(meta_rows):
        left, right = table.rows[row_index].cells
        set_cell_shading(left, LIGHTER_BLUE)
        set_cell_shading(right, WHITE)
        for cell in (left, right):
            set_cell_margins(cell, top=90, bottom=90, start=140, end=140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_inline(p1, label, size=9.5, color=BLUE, bold=True)
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_inline(p2, value, size=10, color=INK)

    callout = doc.add_paragraph(style="Callout")
    callout.paragraph_format.space_before = Pt(22)
    callout.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(callout, LIGHTER_BLUE)
    set_paragraph_border(callout, left={"size": 16, "space": 4, "color": BLUE})
    add_inline(
        callout,
        "核心定位：先建设一条可运营、可追溯、可结算的视频数据生产流水线；"
        "一期不建设大而全的数据交易平台或模型训练平台。",
        size=10.5,
        color=NAVY,
        bold=True,
    )

    doc.add_page_break()


def add_navigation(doc: Document) -> None:
    p = doc.add_paragraph("阅读导航", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    nav = [
        ("业务范围与需求", "第 0-6 章：现状、范围、权限、一期/二期需求和状态机"),
        ("架构与技术选型", "第 7-8 章：推荐架构、云平台、小程序、AI 和媒体 Worker 候选"),
        ("容量与工程接口", "第 9-13 章：1000 用户容量、存储、数据模型、API 和 AI 协议"),
        ("安全与运行保障", "第 14-15 章：安全、隐私、环境、监控、备份和发布"),
        ("实施与决策", "第 16-19 章：人天对齐、验收、待确认事项和最终推荐"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "阅读路径"
    table.rows[0].cells[1].text = "内容"
    for label, content in nav:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = content
    format_table(table, headers=["阅读路径", "内容"])


def choose_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    cols = len(headers)
    if cols == 2:
        if headers[0] in {"角色", "组件", "API 组", "领域", "方案", "工作包"}:
            return [2200, 7160]
        return [1900, 7460]
    if cols == 3:
        if headers[0] == "ID":
            return [1250, 2500, 5610]
        if headers[-1] in {"基准人天", "建议"}:
            return [2100, 5600, 1660]
        return [1800, 3100, 4460]
    if cols == 4:
        return [1800, 2520, 3140, 1900]
    if cols == 5:
        return [1500, 1500, 2000, 2100, 2260]
    if cols == 6:
        return [1250, 1150, 1450, 1650, 2000, 1860]
    return [9360 // cols] * cols


def format_table(table, headers: list[str]) -> None:
    rows_data = [[cell.text for cell in row.cells] for row in table.rows[1:]]
    widths = choose_widths(headers, rows_data)
    set_table_geometry(table, widths, indent_dxa=120)
    set_table_borders(table, color=MID_GRAY, size=4)
    mark_repeat_header(table.rows[0])
    col_count = len(headers)
    font_size = 9.5 if col_count <= 3 else (9 if col_count == 4 else 8.5)

    for r_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for c_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            if r_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif r_index % 2 == 0:
                set_cell_shading(cell, "FAFBFD")
            else:
                set_cell_shading(cell, WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.12
                if c_index == 0 or (
                    c_index == col_count - 1
                    and headers[c_index] in {"基准人天", "规划值", "建议"}
                ):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        color=(NAVY if r_index == 0 else INK),
                        bold=(r_index == 0),
                    )
    after = table._tbl.getnext()
    if after is None or after.tag != qn("w:p"):
        table._tbl.addnext(OxmlElement("w:p"))


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    def split_row(line: str) -> list[str]:
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    headers = split_row(lines[0])
    data_lines = lines[2:]
    rows = [split_row(line) for line in data_lines]
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value.replace("`", "")
    format_table(table, headers)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Callout")
    set_paragraph_shading(p, LIGHTER_BLUE)
    set_paragraph_border(p, left={"size": 14, "space": 4, "color": BLUE})
    add_inline(p, text, size=10.5, color=NAVY)


def add_diagram(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Diagram")
    set_paragraph_shading(p, LIGHTER_BLUE)
    set_paragraph_border(p, left={"size": 12, "space": 4, "color": DARK_BLUE})
    add_inline(p, text, size=9.5, color=NAVY, bold=True)


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in stripped.split("|"))


def parse_markdown(doc: Document, markdown: str, bullet_num: int, decimal_num: int) -> None:
    lines = markdown.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("# 0."))
    index = start
    page_break_headings = set()
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            paragraph = doc.add_paragraph(text, style=f"Heading {level}")
            if level == 1 and text in page_break_headings:
                paragraph.paragraph_format.page_break_before = True
            index += 1
            continue

        if stripped.startswith("> "):
            add_callout(doc, stripped[2:].strip())
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph()
            apply_num(p, bullet_num)
            add_inline(p, bullet_match.group(1))
            index += 1
            continue

        decimal_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if decimal_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.34)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, f"{decimal_match.group(1)}. {decimal_match.group(2)}")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if re.match(r"^(#{1,3})\s+", candidate):
                break
            if candidate.startswith("> ") or candidate.startswith("|"):
                break
            if re.match(r"^-\s+", candidate) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        text = " ".join(paragraph_lines)
        if "→" in text and len(text) < 320:
            add_diagram(doc, text)
        else:
            p = doc.add_paragraph()
            add_inline(p, text)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_sections(doc)
    configure_styles(doc)
    bullet_num = add_numbering(doc, kind="bullet")
    decimal_num = add_numbering(doc, kind="decimal")
    add_cover(doc)
    add_navigation(doc)
    parse_markdown(doc, markdown, bullet_num, decimal_num)

    doc.core_properties.title = "具身智能视频数据平台 - 整体需求清单与技术实现方案"
    doc.core_properties.subject = "需求评审、技术选型、开发拆分与一期验收"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "具身智能, 视频数据, 微信小程序, Cloudflare, D1, R2"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
